#!/usr/bin/env python3
"""extract-intake.py -- canonical client-document extraction for Q investigations.

Pairs with ingest-client-document.sh. That script registers each original file as
an EV-NNNN item with SHA-256 + chain of custody. THIS script does the verbatim
content extraction + full OCR that the evidence protocol requires, deterministically
(no network, no LLM -- same input always yields the same output).

For every file in a case's investigation/intake/, writes:
  investigation/evidence/extracted/<stem>/
    text.md        -- verbatim text (PDF text layer / DOCX paragraphs+tables / DOC via textutil)
    images/        -- embedded images (DOCX/DOC) and full-page renders (PDF)
    ocr.md         -- tesseract OCR (multi-PSM, best-by-alnum) of every image/page
    manifest.json  -- per-file extraction stats for verification

Handlers: .pdf .docx .doc .png .jpg .jpeg .tif .tiff
Requires on PATH: pdftotext, pdftoppm, tesseract. Python: python-docx + Pillow
(only needed when DOCX/DOC/image inputs are present; imported lazily).

Usage:
  python3 extract-intake.py --case <case-folder> [--intake <dir>]
    --case    investigations/<case-folder>; defaults to <workspace>/.active-case
    --intake  override the intake dir (default: <case>/investigation/intake)

Exit codes: 0 all files extracted; 1 setup error (no case/intake/files);
            2 one or more files failed extraction (others still processed).
"""

import argparse
import json
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path

EXTRACTOR_NAME = "extract-intake.py"

# workspace = workspace (script lives in skills/osint/scripts/)
WORKSPACE = Path(__file__).resolve().parents[3]

PSM_MODES = ["3", "6", "11"]
PSM_LABELS = {
    "3": "PSM 3 (auto)",
    "6": "PSM 6 (uniform block)",
    "11": "PSM 11 (sparse text)",
    "up4x": "upscale 4x + binarize + PSM 11",
}
RENDER_DPI = "200"
ALNUM_RE = re.compile(r"[A-Za-z0-9]")
TRAILING_NUM_RE = re.compile(r"(\d+)(?=\D*$)")


# ---------- small shared helpers ----------

def run_cmd(cmd: list) -> str:
    result = subprocess.run(cmd, capture_output=True, text=True, check=False)
    return result.stdout


def alnum_count(text: str) -> int:
    return len(ALNUM_RE.findall(text))


def write_text_md(out_dir: Path, src: Path, text: str, note: str) -> None:
    body = [f"# {src.name}", "", f"_Extraction: {note}_", "",
            "## Verbatim text", "", text.rstrip(), ""]
    (out_dir / "text.md").write_text("\n".join(body), encoding="utf-8")


def sort_pages(paths: list) -> list:
    def key(p: Path) -> int:
        match = TRAILING_NUM_RE.search(p.stem)
        return int(match.group(1)) if match else 0
    return sorted(paths, key=key)


# ---------- OCR (multi-PSM, deterministic) ----------

def tesseract_psm(image: Path, psm: str) -> str:
    result = subprocess.run(
        ["tesseract", str(image), "-", "-l", "eng", "--psm", psm],
        capture_output=True, text=True, check=False,
    )
    return result.stdout.rstrip() if result.returncode == 0 else ""


def upscale_and_binarize(image: Path):
    from PIL import Image
    try:
        img = Image.open(image)
    except Exception:
        return None
    if img.width < 100 or img.height < 100:
        return None
    upscaled = img.resize((img.width * 4, img.height * 4), Image.LANCZOS).convert("L")
    binarized = upscaled.point(lambda p: 255 if p > 140 else 0)
    out_path = image.with_suffix(".up4x.png")
    binarized.save(out_path)
    return out_path


def ocr_image(image: Path):
    """Return (best_text, best_label). Tries 3 PSM modes, upscales if all empty."""
    outputs = {psm: tesseract_psm(image, psm) for psm in PSM_MODES}
    if all(alnum_count(outputs[p]) == 0 for p in PSM_MODES):
        upscaled = upscale_and_binarize(image)
        outputs["up4x"] = tesseract_psm(upscaled, "11") if upscaled else ""
    else:
        outputs["up4x"] = ""
    order = PSM_MODES + ["up4x"]
    best = max(order, key=lambda p: (alnum_count(outputs[p]), -order.index(p)))
    return outputs[best], best


def ocr_images(out_dir: Path, src: Path, images: list, header: str) -> int:
    if not images:
        return 0
    lines = [f"# {src.name} -- OCR ({header})", ""]
    total_chars = 0
    for img in images:
        best_text, best_label = ocr_image(img)
        total_chars += len(best_text)
        lines += [f"## {img.name}  ({PSM_LABELS[best_label]})", "",
                  "```", best_text or "(no text detected)", "```", ""]
    (out_dir / "ocr.md").write_text("\n".join(lines), encoding="utf-8")
    return total_chars


# ---------- DOCX helpers ----------

def docx_text(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    lines = ["## Paragraphs", ""]
    for para in doc.paragraphs:
        if para.text.strip():
            lines += [para.text, ""]
    table_no = 0
    for table in doc.tables:
        table_no += 1
        lines += [f"## Table {table_no}", ""]
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")
    return "\n".join(lines)


def docx_images(path: Path, images_dir: Path) -> list:
    from docx import Document
    images_dir.mkdir(parents=True, exist_ok=True)
    doc = Document(str(path))
    saved = []
    for rel in doc.part.related_parts.values():
        content_type = getattr(rel, "content_type", "") or ""
        if not content_type.startswith("image/"):
            continue
        ext = content_type.split("/")[-1].split("+")[0]
        if ext == "jpeg":
            ext = "jpg"
        img_path = images_dir / f"image_{len(saved) + 1:03d}.{ext}"
        img_path.write_bytes(rel.blob)
        saved.append(img_path)
    return saved


# ---------- per-type handlers ----------

def extract_pdf(path: Path, out: Path) -> dict:
    out.mkdir(parents=True, exist_ok=True)
    images_dir = out / "images"
    images_dir.mkdir(exist_ok=True)

    text = run_cmd(["pdftotext", "-layout", str(path), "-"])
    write_text_md(out, path, text, "PDF text layer via pdftotext -layout")

    subprocess.run(["pdftoppm", "-png", "-r", RENDER_DPI, str(path), str(images_dir / "page")],
                   check=False)
    pages = sort_pages(list(images_dir.glob("page*.png")))
    ocr_images(out, path, pages, f"full-page renders, {len(pages)} pages")
    return {"text_chars": len(text), "images": len(pages)}


def extract_docx(path: Path, out: Path) -> dict:
    out.mkdir(parents=True, exist_ok=True)
    text = docx_text(path)
    write_text_md(out, path, text, "DOCX paragraphs + tables via python-docx")
    images = docx_images(path, out / "images")
    ocr_images(out, path, images, "DOCX embedded images")
    return {"text_chars": len(text), "images": len(images)}


def extract_doc(path: Path, out: Path) -> dict:
    out.mkdir(parents=True, exist_ok=True)
    text = run_cmd(["textutil", "-convert", "txt", "-stdout", str(path)])
    write_text_md(out, path, text, "legacy .doc text via textutil")

    converted = out / f"{path.stem}.converted.docx"
    subprocess.run(["textutil", "-convert", "docx", "-output", str(converted), str(path)],
                   check=False)
    images = docx_images(converted, out / "images") if converted.exists() else []
    ocr_images(out, path, images, "DOC embedded images (textutil -> docx)")
    return {"text_chars": len(text), "images": len(images)}


def extract_image(path: Path, out: Path) -> dict:
    out.mkdir(parents=True, exist_ok=True)
    images_dir = out / "images"
    images_dir.mkdir(exist_ok=True)
    dest = images_dir / path.name
    shutil.copy2(path, dest)
    write_text_md(out, path, "(binary image -- OCR only, see ocr.md)", "image file")
    ocr_images(out, path, [dest], "image OCR")
    return {"text_chars": 0, "images": 1}


def extract_csv(path: Path, out: Path) -> dict:
    import csv
    out.mkdir(parents=True, exist_ok=True)
    lines = [f"# {path.name}", "", "_Extraction: CSV rows as structured text (csv module)_", ""]
    row_count = 0
    headers = []
    try:
        with path.open(encoding="utf-8", errors="replace", newline="") as fh:
            reader = csv.DictReader(fh)
            headers = reader.fieldnames or []
            for i, row in enumerate(reader):
                non_empty = {k: v for k, v in row.items() if k and v and v.strip()}
                if not non_empty:
                    continue
                lines.append(f"## Row {i + 1}")
                lines.append("")
                for k, v in non_empty.items():
                    lines.append(f"**{k}:** {v.strip()}")
                lines.append("")
                row_count += 1
    except Exception as exc:
        lines.append(f"_(CSV parse error: {exc})_")
    text = "\n".join(lines)
    (out / "text.md").write_text(text, encoding="utf-8")
    # No images in CSV; write empty ocr.md for consistency
    (out / "ocr.md").write_text(f"# {path.name} -- OCR\n\n_(CSV file -- no images to OCR)_\n",
                                encoding="utf-8")
    return {"text_chars": len(text), "images": 0, "rows": row_count, "columns": len(headers)}


def extract_eml(path: Path, out: Path) -> dict:
    """Extract an .eml email: headers, body (plain text, HTML fallback),
    inline images (OCR'd via the shared multi-PSM pipeline), attachments
    saved verbatim, and any .docx attachment additionally extracted into
    its own subdirectory. Ported 2026-07-21 from case-025/case-028's
    independently-written extract_intake.py -- this capability (email
    extraction) did not exist in this module until now."""
    import email
    import email.policy

    out.mkdir(parents=True, exist_ok=True)
    images_dir = out / "images"
    attachments_dir = out / "attachments"
    images_dir.mkdir(exist_ok=True)
    attachments_dir.mkdir(exist_ok=True)

    msg = email.message_from_bytes(path.read_bytes(), policy=email.policy.default)

    lines = [f"# {path.name}", "", "## Headers", "",
             f"- **From:** {msg.get('From', '(none)')}",
             f"- **To:** {msg.get('To', '(none)')}",
             f"- **CC:** {msg.get('Cc', '(none)')}",
             f"- **Date:** {msg.get('Date', '(none)')}",
             f"- **Subject:** {msg.get('Subject', '(none)')}",
             "", "## Body", ""]

    body_parts, image_paths, attachment_paths = [], [], []
    for part in msg.walk():
        content_type = part.get_content_type()
        disposition = str(part.get("Content-Disposition", ""))

        if content_type == "text/plain" and "attachment" not in disposition:
            try:
                body_parts.append(part.get_content())
            except Exception as exc:
                body_parts.append(f"(decode error: {exc})")
        elif content_type == "text/html" and "attachment" not in disposition and not body_parts:
            try:
                html = part.get_content()
                body_parts.append(f"[HTML body -- first 4000 chars]\n{html[:4000]}")
            except Exception as exc:
                body_parts.append(f"(HTML decode error: {exc})")
        elif content_type.startswith("image/"):
            ext = content_type.split("/")[-1].split("+")[0]
            ext = "jpg" if ext == "jpeg" else ext
            img_path = images_dir / f"image_{len(image_paths) + 1:03d}.{ext}"
            try:
                img_path.write_bytes(part.get_payload(decode=True))
                image_paths.append(img_path)
            except Exception as exc:
                print(f"  [warn] inline image failed: {exc}", file=sys.stderr)
        elif "attachment" in disposition or content_type not in (
            "text/plain", "text/html", "multipart/mixed",
            "multipart/alternative", "multipart/related", "multipart/signed",
        ):
            filename = part.get_filename()
            if filename:
                att_path = attachments_dir / filename
                try:
                    att_path.write_bytes(part.get_payload(decode=True))
                    attachment_paths.append(att_path)
                except Exception as exc:
                    print(f"  [warn] attachment {filename} failed: {exc}", file=sys.stderr)

    lines.extend(body_parts if body_parts else ["(no plain-text body found)"])
    lines.append("")
    if attachment_paths:
        lines += ["## Attachments", ""]
        lines += [f"- `{a.name}` -> `attachments/{a.name}`" for a in attachment_paths]
        lines.append("")
    if image_paths:
        lines += ["## Inline Images", ""]
        lines += [f"- `{i.name}` -> `images/{i.name}`" for i in image_paths]
        lines.append("")
    (out / "text.md").write_text("\n".join(lines), encoding="utf-8")

    ocr_images(out, path, image_paths, "EML inline images")

    for att in attachment_paths:
        if att.suffix.lower() == ".docx":
            try:
                extract_docx(att, attachments_dir / f"{att.stem}_extracted")
            except Exception as exc:
                print(f"  [warn] nested docx extract failed for {att.name}: {exc}", file=sys.stderr)

    return {"text_chars": sum(len(b) for b in body_parts), "images": len(image_paths),
            "attachments": len(attachment_paths)}


def extract_md(path: Path, out: Path) -> dict:
    out.mkdir(parents=True, exist_ok=True)
    text = path.read_text(encoding="utf-8", errors="replace")
    write_text_md(out, path, text, "Markdown file, copied verbatim")
    (out / "ocr.md").write_text(f"# {path.name} -- OCR\n\n_(Markdown file -- no images to OCR)_\n",
                                encoding="utf-8")
    return {"text_chars": len(text), "images": 0}


HANDLERS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_doc,
    ".csv": extract_csv,
    ".eml": extract_eml,
    ".md": extract_md,
    ".png": extract_image,
    ".jpg": extract_image,
    ".jpeg": extract_image,
    ".tif": extract_image,
    ".tiff": extract_image,
}


# ---------- orchestration ----------

def read_active_case() -> str:
    pointer = WORKSPACE / ".active-case"
    if not pointer.is_file():
        return ""
    return pointer.read_text().strip()


def resolve_paths(args):
    case = args.case or read_active_case()
    if not case:
        print("ERROR: no --case given and no .active-case set", file=sys.stderr)
        sys.exit(1)
    case_dir = WORKSPACE / "investigations" / case
    if not case_dir.is_dir():
        print(f"ERROR: case dir not found: {case_dir}", file=sys.stderr)
        sys.exit(1)
    intake = Path(args.intake) if args.intake else case_dir / "investigation" / "intake"
    if not intake.is_dir():
        print(f"ERROR: intake dir not found: {intake}", file=sys.stderr)
        sys.exit(1)
    extracted = case_dir / "investigation" / "evidence" / "extracted"
    extracted.mkdir(parents=True, exist_ok=True)
    return case, intake, extracted


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract verbatim text + full OCR from intake files.")
    parser.add_argument("--case", help="case folder name under investigations/")
    parser.add_argument("--intake", help="override intake dir")
    args = parser.parse_args()

    case, intake, extracted = resolve_paths(args)
    files = sorted(p for p in intake.iterdir() if p.is_file())
    if not files:
        print(f"ERROR: no files in {intake}", file=sys.stderr)
        return 1

    print(f"Case: {case}")
    print(f"Intake: {intake}  ({len(files)} files)\n")
    print(f"{'FILE':<52} {'TEXT_CH':>8} {'IMGS':>5} {'OCR_KB':>7} {'STATUS':>7}")

    failures = []
    for src in files:
        handler = HANDLERS.get(src.suffix.lower())
        if handler is None:
            print(f"{src.name[:52]:<52} {'-':>8} {'-':>5} {'-':>7} {'SKIP':>7}")
            continue
        out = extracted / src.stem
        try:
            stats = handler(src, out)
            ocr_kb = (out / "ocr.md").stat().st_size / 1024 if (out / "ocr.md").exists() else 0
            stats.update({"file": src.name, "ext": src.suffix.lower(),
                          "ocr_bytes": int(ocr_kb * 1024), "status": "OK",
                          "extractor": EXTRACTOR_NAME,
                          "extracted_at": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")})
            (out / "manifest.json").write_text(json.dumps(stats, indent=2), encoding="utf-8")
            print(f"{src.name[:52]:<52} {stats['text_chars']:>8} {stats['images']:>5} "
                  f"{ocr_kb:>7.1f} {'OK':>7}")
        except Exception as exc:  # noqa: BLE001 -- report and continue, preserve partial evidence
            failures.append((src.name, str(exc)))
            print(f"{src.name[:52]:<52} {'-':>8} {'-':>5} {'-':>7} {'FAIL':>7}")

    print(f"\nExtracted -> {extracted}")
    if failures:
        print("\nFAILURES:", file=sys.stderr)
        for name, err in failures:
            print(f"  {name}: {err}", file=sys.stderr)
        return 2
    return 0


if __name__ == "__main__":
    sys.exit(main())
